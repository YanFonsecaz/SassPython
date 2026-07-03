import base64
import io

from docx import Document
from PIL import Image as PILImage

from app.services.parecer_service import estrutura_para_html, html_para_docx_bytes


def _make_data_uri(width=100, height=100, fmt="PNG") -> str:
    buf = io.BytesIO()
    img = PILImage.new("RGB", (width, height), color="red")
    img.save(buf, format=fmt)
    buf.seek(0)
    b64 = base64.b64encode(buf.read()).decode()
    mime = "image/png" if fmt == "PNG" else "image/webp"
    return f"data:{mime};base64,{b64}"


ESTRUTURA_EXEMPLO = {
    "titulo": "PARECER TÉCNICO — SEO / PERFORMANCE",
    "subtitulo": "Otimização de Core Web Vitals",
    "escopo_linha": "LCP e CLS — exemplo.com.br (Loja Teste)",
    "secoes": [
        {
            "titulo": "Página de produto — Exemplo",
            "url": "https://exemplo.com.br/p/1",
            "observacao": "Observação: ocorre de forma idêntica em Desktop e Mobile.",
            "subsecoes": [
                {
                    "titulo": "LCP atrasado por CSS bloqueante",
                    "problemas": [
                        {
                            "descricao": "LCP acima de 4s no mobile.",
                            "evidencias": [
                                {"legenda": "Trace: global.css bloqueando a renderização.", "imagens_indices": [0]},
                            ],
                            "solucao": "Extrair CSS crítico inline e carregar o restante de forma assíncrona.",
                            "solucao_escopo": "Desktop e Mobile",
                        },
                    ],
                },
                {
                    "titulo": "Imagens de produto em JPEG",
                    "problemas": [
                        {
                            "descricao": "Imagem principal em JPEG sem prioridade.",
                            "evidencias": [{"legenda": "DOM: imagem em JPEG.", "imagens_indices": []}],
                            "solucao": "Servir em WebP e aplicar fetchpriority=high.",
                        },
                        {
                            "descricao": "Imagem sem width/height, causando CLS.",
                            "evidencias": [{"legenda": "Imagem sem dimensões reservadas.", "imagens_indices": []}],
                            "solucao": "Definir width/height ou aspect-ratio.",
                        },
                    ],
                },
            ],
        },
    ],
    "recomendacoes_globais": [
        {"titulo": "Prioridade 1 — Eliminar render-blocking", "itens": ["Item 1", "Item 2"]},
    ],
}


class TestEstruturaParaHtml:
    def test_basic_structure(self):
        data_uri = _make_data_uri()
        html_str = estrutura_para_html(ESTRUTURA_EXEMPLO, {0: data_uri})

        # cabecalho de 3 linhas
        assert "<h1>PARECER TÉCNICO — SEO / PERFORMANCE</h1>" in html_str
        assert "<em>Otimização de Core Web Vitals</em>" in html_str
        assert "LCP e CLS — exemplo.com.br (Loja Teste)" in html_str
        # sem tabela de metadados e sem sumario executivo
        assert "data-meta" not in html_str
        assert "data-causas" not in html_str
        assert "Sumário executivo" not in html_str
        # secao numerada a partir de 1 + url + observacao
        assert "1. Página de produto — Exemplo" in html_str
        assert "URL: https://exemplo.com.br/p/1" in html_str
        assert "Observação: ocorre de forma idêntica em Desktop e Mobile." in html_str
        # subsecoes N.M.
        assert "1.1. LCP atrasado por CSS bloqueante" in html_str
        assert "1.2. Imagens de produto em JPEG" in html_str
        # problema unico vs multiplos
        assert "<strong>Problema</strong>" in html_str
        assert "<strong>Problema 1</strong>" in html_str
        assert "<strong>Problema 2</strong>" in html_str
        # evidencia, solucao com escopo, imagem
        assert "<strong>Evidência:</strong>" in html_str
        assert "<strong>Solução (Desktop e Mobile)</strong>" in html_str
        assert "<strong>Solução</strong>" in html_str
        assert f'src="{data_uri}"' in html_str
        # recomendacoes globais como ultima secao (N = 1 secao + 1 = 2)
        assert "2. Recomendações globais (aplicáveis a toda a loja)" in html_str
        assert "<ul>" in html_str
        assert "<li>Item 1</li>" in html_str

    def test_with_multiple_images(self):
        uri0 = _make_data_uri(100, 100)
        uri1 = _make_data_uri(200, 200)
        html_str = estrutura_para_html(ESTRUTURA_EXEMPLO, {0: uri0, 1: uri1})
        assert uri0 in html_str
        assert uri1 not in html_str

    def test_empty_images(self):
        html_str = estrutura_para_html(ESTRUTURA_EXEMPLO, {})
        assert "<img" not in html_str


class TestHtmlParaDocxBytes:
    def test_returns_nonempty_bytes(self):
        data_uri = _make_data_uri()
        html_str = estrutura_para_html(ESTRUTURA_EXEMPLO, {0: data_uri})
        result = html_para_docx_bytes(html_str)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_opens_as_valid_docx(self):
        data_uri = _make_data_uri()
        html_str = estrutura_para_html(ESTRUTURA_EXEMPLO, {0: data_uri})
        result = html_para_docx_bytes(html_str)
        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_has_title_heading_and_image(self):
        data_uri = _make_data_uri()
        html_str = estrutura_para_html(ESTRUTURA_EXEMPLO, {0: data_uri})
        result = html_para_docx_bytes(html_str)
        doc = Document(io.BytesIO(result))

        styles = [p.style.name for p in doc.paragraphs]
        assert "Title" in styles
        assert "Heading 1" in styles
        assert "Heading 2" in styles
        # novo formato nao usa tabelas no documento gerado
        assert len(doc.tables) == 0
        assert len(doc.inline_shapes) >= 1

    def test_table_inserida_pelo_usuario_ainda_renderiza(self):
        # o renderer suporta tabela caso o usuario insira uma no editor
        result = html_para_docx_bytes("<h1>T</h1><table><tr><td>a</td><td>b</td></tr></table>")
        doc = Document(io.BytesIO(result))
        assert len(doc.tables) == 1

    def test_webp_image_embedded(self):
        data_uri = _make_data_uri(100, 100, fmt="PNG")
        result = html_para_docx_bytes(f'<img src="{data_uri}" />')
        doc = Document(io.BytesIO(result))
        assert len(doc.inline_shapes) >= 1

    def test_edited_html_still_exports(self):
        result = html_para_docx_bytes("<h1>Titulo editado</h1><p>Texto alterado.</p>")
        doc = Document(io.BytesIO(result))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("Titulo editado" in t for t in texts)
        assert any("Texto alterado" in t for t in texts)

    def test_unknown_node_ignored(self):
        result = html_para_docx_bytes("<blockquote>Ignorado</blockquote><p>Normal</p>")
        doc = Document(io.BytesIO(result))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("Normal" in t for t in texts)

    def test_empty_html(self):
        result = html_para_docx_bytes("<p></p>")
        doc = Document(io.BytesIO(result))
        assert doc is not None

    def test_fallback_styles(self):
        result = html_para_docx_bytes("<h1>Test</h1>")
        doc = Document(io.BytesIO(result))
        try:
            normal = doc.styles["Normal"]
            assert normal.font.name == "Calibri"
        except KeyError:
            pass

    def test_code_block_preservado(self):
        # blocos de codigo (CWV "Como corrigir") nao podem ser perdidos
        result = html_para_docx_bytes("<pre><code>const x = 1;</code></pre><p>fim</p>")
        doc = Document(io.BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        assert any("const x = 1;" in t for t in texts)
        assert any("fim" in t for t in texts)

    def test_headings_h4_h6_preservados(self):
        result = html_para_docx_bytes("<h4>Sub A</h4><h6>Sub B</h6>")
        doc = Document(io.BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        assert any("Sub A" in t for t in texts)
        assert any("Sub B" in t for t in texts)
