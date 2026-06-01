import io

from docx import Document

from app.services.cwv_export import (
    problema_para_html,
    relatorio_para_html,
    slugify_titulo,
)
from app.services.parecer_service import html_para_docx_bytes

PROBLEMA = {
    "titulo": "Scripts bloqueando renderizacao no <head>",
    "severidade": 5,
    "metricas_afetadas": ["LCP", "FCP"],
    "contexto_especifico": {
        "display_value": "Est savings of 1,200 ms",
        "savings_ms": 1200,
        "description": "Requests are blocking render.",
        "items": [
            {"url": "https://site.com/a.css", "wastedBytes": 1600, "totalBytes": 1600},
            {"url": "https://site.com/b.css", "wastedMs": 455, "totalBytes": 1200},
        ],
    },
    "documentacao_md": "## Problema\nScripts no head.\n\n## Solucao\n1. Mova\n\n```\n<Script src=\"x.js\" />\n```\n",
}

ANALISE = {
    "url_canonica": "https://site.com/produto",
    "plataforma_detectada": "Next.js",
    "estrategia": "mobile",
    "criado_em": "2026-06-01T10:00:00",
    "score_performance": 60,
    "lcp_ms": 8000,
    "cls": 0.1,
    "inp_ms": 53,
    "fcp_ms": 1200,
    "usuario_id": "u1",
}


class TestProblemaParaHtml:
    def test_inclui_titulo_recursos_e_como_corrigir(self):
        html = problema_para_html(PROBLEMA)
        assert "Scripts bloqueando" in html
        assert "<table" in html  # recursos como HTML (nao markdown)
        assert "Como corrigir" in html

    def test_docx_tem_tabela_de_recursos_e_codigo(self):
        result = html_para_docx_bytes(problema_para_html(PROBLEMA))
        doc = Document(io.BytesIO(result))
        # a tabela de recursos deve existir de verdade (nao virar texto cru)
        assert len(doc.tables) >= 1
        cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
        assert any("a.css" in c for c in cells)
        # bloco de codigo do "Como corrigir" preservado
        full = "\n".join(p.text for p in doc.paragraphs)
        assert "<Script" in full
        # nao pode sobrar markdown de tabela cru
        assert "| Recurso" not in full


class TestRelatorioParaHtml:
    def test_relatorio_tem_metricas_sumario_e_problemas(self):
        result = html_para_docx_bytes(relatorio_para_html(ANALISE, [PROBLEMA, PROBLEMA]))
        doc = Document(io.BytesIO(result))
        # metricas + sumario + 2 tabelas de recursos
        assert len(doc.tables) >= 3
        cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
        assert any("8000 ms" in c for c in cells)  # LCP formatado em ms
        headings = [p.text for p in doc.paragraphs]
        assert any("Sumario" in h for h in headings)
        assert any("1. Scripts" in h for h in headings)


def test_slugify_titulo():
    assert slugify_titulo("Scripts Bloqueando!! no <head>") == "scripts-bloqueando-no-head"
    assert slugify_titulo("") == "cwv-problema"
