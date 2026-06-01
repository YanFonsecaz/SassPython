import base64
import html
import io
import logging
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from PIL import Image
from selectolax.parser import HTMLParser

TEMPLATE_PATH = Path(__file__).parent / "templates" / "parecer_template.docx"
LARGURA_UTIL = Inches(6.3)

logger = logging.getLogger(__name__)


def estrutura_para_html(estrutura: dict, imagens_por_indice: dict[int, str]) -> str:
    """Renderiza o parecer no padrao do documento de referencia (Imecap):
    cabecalho de 3 linhas -> secoes por pagina (N.) -> subsecoes (N.M.) com
    Problema/Evidencia(s)/Solucao -> "Recomendacoes globais" como ultima secao.
    Sem tabela de metadados e sem sumario executivo.
    """
    parts: list[str] = []

    # Cabecalho: titulo + subtitulo + linha de escopo
    parts.append(f"<h1>{html.escape(estrutura.get('titulo', ''))}</h1>")
    subtitulo = estrutura.get("subtitulo", "")
    if subtitulo:
        parts.append(f"<p><em>{html.escape(subtitulo)}</em></p>")
    escopo = estrutura.get("escopo_linha", "")
    if escopo:
        parts.append(f"<p>{html.escape(escopo)}</p>")

    secoes = estrutura.get("secoes", [])
    for i, secao in enumerate(secoes, start=1):
        parts.append(f"<h2>{i}. {html.escape(secao.get('titulo', ''))}</h2>")
        url = secao.get("url")
        if url:
            parts.append(f"<p>URL: {html.escape(url)}</p>")
        obs = secao.get("observacao")
        if obs:
            parts.append(f"<p><em>{html.escape(obs)}</em></p>")
        for j, sub in enumerate(secao.get("subsecoes", []), start=1):
            parts.append(f"<h3>{i}.{j}. {html.escape(sub.get('titulo', ''))}</h3>")
            problemas = sub.get("problemas", [])
            multi = len(problemas) > 1
            for k, prob in enumerate(problemas, start=1):
                rotulo = f"Problema {k}" if multi else "Problema"
                parts.append(f"<p><strong>{rotulo}</strong> {html.escape(prob.get('descricao', ''))}</p>")
                for ev in prob.get("evidencias", []):
                    parts.append(f"<p><strong>Evidência:</strong> {html.escape(ev.get('legenda', ''))}</p>")
                    for idx in ev.get("imagens_indices", []):
                        data_uri = imagens_por_indice.get(idx, "")
                        if data_uri:
                            parts.append(f'<img src="{data_uri}" />')
                escopo_sol = prob.get("solucao_escopo")
                sol_label = f"Solução ({escopo_sol})" if escopo_sol else "Solução"
                parts.append(f"<p><strong>{sol_label}</strong> {html.escape(prob.get('solucao', ''))}</p>")

    recomendacoes = estrutura.get("recomendacoes_globais", [])
    if recomendacoes:
        n = len(secoes) + 1
        parts.append(f"<h2>{n}. Recomendações globais (aplicáveis a toda a loja)</h2>")
        for prioridade in recomendacoes:
            parts.append(f"<p><strong>{html.escape(prioridade.get('titulo', ''))}</strong></p>")
            itens = prioridade.get("itens", [])
            if itens:
                parts.append("<ul>")
                for item in itens:
                    parts.append(f"<li>{html.escape(item)}</li>")
                parts.append("</ul>")

    return "\n".join(parts)


def html_para_docx_bytes(html_str: str) -> bytes:
    doc = Document(str(TEMPLATE_PATH)) if TEMPLATE_PATH.exists() else Document()
    _garantir_estilos(doc)
    tree = HTMLParser(html_str)
    body = tree.body or tree.root
    for node in body.iter(include_text=False):
        _render_node(doc, node)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _garantir_estilos(doc: Document) -> None:
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
    except KeyError:
        pass
    try:
        h1 = doc.styles["Title"]
        h1.font.color.rgb = RGBColor(0x5C, 0x52, 0x49)
    except (KeyError, AttributeError):
        pass


def _node_text(node) -> str:
    t = node.text()
    return t if t else ""


def _render_node(doc: Document, node) -> None:
    tag = node.tag.lower() if node.tag else ""
    if tag == "h1":
        _add_paragraph(doc, _node_text(node), style="Title")
    elif tag == "h2":
        _add_paragraph(doc, _node_text(node), style="Heading 1")
    elif tag == "h3":
        _add_paragraph(doc, _node_text(node), style="Heading 2")
    elif tag == "p":
        _render_paragraph(doc, node)
    elif tag == "table":
        _render_table(doc, node)
    elif tag in ("ul", "ol"):
        _render_list(doc, node, ordered=(tag == "ol"))
    elif tag == "img":
        src = node.attributes.get("src", "")
        if src.startswith("data:"):
            _add_imagem(doc, src)
    elif tag == "hr":
        _add_paragraph(doc, "—", style="Normal")
    else:
        for child in node.iter(include_text=False):
            _render_node(doc, child)


def _add_paragraph(doc: Document, text: str, style: str = "Normal") -> None:
    if text.strip():
        doc.add_paragraph(text.strip(), style=style)


def _render_paragraph(doc: Document, node) -> None:
    p = doc.add_paragraph(style="Normal")
    _render_inline(p, node)


def _render_inline(paragraph, node) -> None:
    for child in node.iter(include_text=True):
        ctag = (child.tag or "").lower()
        text = child.text() or ""
        if not text:
            continue
        run = paragraph.add_run(text)
        if ctag in ("strong", "b"):
            run.bold = True
        elif ctag in ("em", "i"):
            run.italic = True


def _render_table(doc: Document, node) -> None:
    data_type = node.attributes.get("data-meta") or node.attributes.get("data-causas")
    rows = []
    for tr in node.css("tr"):
        cells = [td.text().strip() for td in tr.iter() if td.tag == "td" or td.tag == "th"]
        rows.append(cells)
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols, style="Table Grid")
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    if i == 0 and (data_type == "data-causas" or data_type):
                        run.bold = True


def _render_list(doc: Document, node, ordered: bool = False) -> None:
    for li in node.css("li"):
        text = li.text().strip()
        if text:
            style = "List Number" if ordered else "List Bullet"
            doc.add_paragraph(text, style=style)


def _add_imagem(doc: Document, data_uri: str) -> None:
    try:
        if "," not in data_uri:
            return
        _header, b64 = data_uri.split(",", 1)
        raw = base64.b64decode(b64)
        img = Image.open(io.BytesIO(raw))
        if getattr(img, "is_animated", False):
            img.seek(0)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        png = io.BytesIO()
        img.save(png, "PNG")
        png.seek(0)
        largura = min(LARGURA_UTIL, Inches(img.width / 96))
        doc.add_picture(png, width=largura)
    except Exception:
        logger.warning("Falha ao embutir imagem no docx", exc_info=True)
